"""
OpenAI Service
Handles OpenAI API interactions for monitoring plan extraction

Enhancements:
- Flexible, hallucination-safe prompt tailored to CORSIA EMP
- Multi-format ingestion (.xlsx/.xls, .csv, .pdf, .docx, images)
- Sheet/page-aware content formatting to support simple provenance
"""

import os
import base64
from typing import Dict, Any, Optional, List
from openai import OpenAI
import json
import PyPDF2
import pandas as pd
from PIL import Image
import io
import pdfplumber
from docx import Document


class OpenAIService:
    """Service for OpenAI API operations"""

    def __init__(self):
        self.client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
        self.model = "gpt-5"  # Using GPT-5 as specified

    def extract_monitoring_plan_info(self, file_path: str, file_extension: str) -> Dict[str, Any]:
        """
        Extract monitoring plan information from uploaded file using GPT-5

        Args:
            file_path: Path to the uploaded file
            file_extension: File extension (pdf, xlsx, csv, png, etc.)

        Returns:
            Extracted monitoring plan information
        """
        print(f"[OPENAI SERVICE] Extracting info from {file_path} ({file_extension})")

        # Extract text/content based on file type
        content = self._extract_content_from_file(file_path, file_extension)

        # Prepare the extraction prompt
        prompt = self._build_extraction_prompt()

        try:
            # Build strict JSON Schema for structured outputs
            schema_block = {
                "type": "json_schema",
                "json_schema": self._build_extraction_schema(),
            }

            # Call GPT-5 with strict schema enforcement
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are an expert in ICAO CORSIA Emissions Monitoring Plans (Annex 16, Vol IV). "
                            "Extract only what is explicitly present. Do not infer or fabricate. If a field is unknown, set it to null and list it in missing_fields. "
                            "Provide concise evidence excerpts and simple provenance identifiers where possible."
                        ),
                    },
                    {
                        "role": "user",
                        "content": f"{prompt}\n\nSOURCE CONTENT (with markers for simple provenance):\n{content}"
                    }
                ],
                reasoning_effort="high",
                response_format=schema_block,
            )

            # Parse the response
            result = json.loads(response.choices[0].message.content)
            print(f"[OPENAI SERVICE] Extraction successful")

            return self._process_extracted_data(result)

        except Exception as e:
            print(f"[OPENAI SERVICE ERROR] Extraction failed: {str(e)}")
            raise Exception(f"Failed to extract monitoring plan information: {str(e)}")

    def _extract_content_from_file(self, file_path: str, file_extension: str) -> str:
        """Extract text content from various file formats and add lightweight provenance markers."""

        try:
            ext = (file_extension or '').lower()
            if ext == 'pdf':
                return self._extract_from_pdf(file_path)
            if ext in ['xlsx', 'xls']:
                return self._extract_from_excel(file_path)
            if ext == 'csv':
                return self._extract_from_csv(file_path)
            if ext in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp']:
                return self._extract_from_image(file_path)
            if ext in ['docx']:
                return self._extract_from_docx(file_path)
            if ext in ['txt', 'text']:
                return self._extract_from_txt(file_path)
            # If we get here, the format is unsupported
                raise ValueError(f"Unsupported file format: {file_extension}")

        except Exception as e:
            print(f"[OPENAI SERVICE ERROR] Content extraction failed: {str(e)}")
            raise

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF. Prefer pdfplumber (layout-aware), fallback to PyPDF2."""
        chunks: List[str] = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages, start=1):
                    page_text = page.extract_text() or ""
                    # Add a page marker for simple provenance
                    chunks.append(f"\n=== PAGE {i} START ===\n{page_text}\n=== PAGE {i} END ===")
        except Exception:
            # Fallback to PyPDF2
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for i, page in enumerate(reader.pages, start=1):
                    page_text = page.extract_text() or ""
                    chunks.append(f"\n=== PAGE {i} START ===\n{page_text}\n=== PAGE {i} END ===")
        return "\n".join(chunks).strip()

    def _extract_from_excel(self, file_path: str) -> str:
        """Extract content from all sheets, preserving sheet names for provenance."""
        sheets = pd.read_excel(file_path, sheet_name=None)  # All sheets
        parts: List[str] = []
        for sheet_name, sheet_df in sheets.items():
            try:
                sheet_df = sheet_df.copy()
                sheet_df.columns = [str(c).strip() for c in sheet_df.columns]
            except Exception:
                pass
            parts.append(f"\n=== SHEET {sheet_name} START ===")
            parts.append(sheet_df.to_string(index=False))
            parts.append(f"=== SHEET {sheet_name} END ===")
        return "\n".join(parts).strip()

    def _extract_from_csv(self, file_path: str) -> str:
        """Extract text from CSV with robust encoding and include a single-sheet marker."""
        try:
            df = pd.read_csv(file_path, encoding='utf-8', encoding_errors='replace')
        except Exception:
            df = pd.read_csv(file_path)
        try:
            df.columns = df.columns.str.strip()
        except Exception:
            pass
        return f"=== SHEET CSV START ===\n{df.to_string(index=False)}\n=== SHEET CSV END ==="

    def _extract_from_image(self, file_path: str) -> str:
        """Extract text from image using GPT-5 Vision"""
        # Read and encode image
        with open(file_path, 'rb') as image_file:
            image_data = base64.b64encode(image_file.read()).decode('utf-8')

        # Determine image format
        file_ext = file_path.split('.')[-1].lower()
        mime_type = f"image/{file_ext if file_ext != 'jpg' else 'jpeg'}"

        # Use GPT-5 with vision to extract text
        # Note: GPT-5 only supports default temperature (1)
        response = self.client.chat.completions.create(
            model=self.model,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Extract all text content from this image. Preserve the structure and formatting as much as possible."
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{mime_type};base64,{image_data}"
                            }
                        }
                    ]
                }
            ]
        )

        return response.choices[0].message.content

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract paragraphs and tables from DOCX with simple section markers."""
        doc = Document(file_path)
        parts: List[str] = []
        parts.append("=== DOCX PARAGRAPHS START ===")
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if text:
                parts.append(text)
        parts.append("=== DOCX PARAGRAPHS END ===")

        # Tables
        if getattr(doc, 'tables', None):
            for t_idx, table in enumerate(doc.tables, start=1):
                parts.append(f"=== DOCX TABLE {t_idx} START ===")
                for row in table.rows:
                    cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                    parts.append(" | ".join(cells))
                parts.append(f"=== DOCX TABLE {t_idx} END ===")
        return "\n".join(parts).strip()

    def _extract_from_txt(self, file_path: str) -> str:
        """Extract plain text with a top-level marker."""
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            data = f.read()
        return f"=== TEXT START ===\n{data}\n=== TEXT END ==="

    def _build_extraction_prompt(self) -> str:
        """Build a flexible, hallucination-safe extraction prompt for GPT-5."""
        return (
            "Extract all relevant information from the CORSIA Emissions Monitoring Plan (EMP). "
            "Return a SINGLE JSON object. Do not include explanations. Rules: "
            "1) Only extract what is explicitly present in the source. 2) If a field is unknown, set null and add a string to missing_fields. \n\n"
            "Extract the following key information:\n"
            "Operator & Regulatory: Operator name; ICAO designator; AOC ID & authority; State of the operator; parent–subsidiary status.\n"
            "Contacts: Name(s), title(s), email(s), phone(s) for accountable manager and CORSIA focal(s).\n"
            "Scope & Period: Reporting year; coverage (international only); aggregation level (state-pair / aerodrome-pair).\n"
            "Fleet & Fuel: Aircraft types and counts used for reporting; fuel types; fuel suppliers if listed.\n"
            "Monitoring Methods (by sub-fleet if applicable): Method A, Method B, Block-off/Block-on, Fuel Uplift, Fuel Allocation (Block Hour); periods of applicability (only 2021–2035); primary data sources (e.g., journey logs/NAVONE, fuel receipts, QAR/ACMS).\n"
            "Calculation Inputs: Fuel density values and sources; CERT usage (yes/no); CERT inputs (e.g., GCD vs actual distances); emission factors if stated.\n"
            "Data Flow & Controls (QA/QC): How data is captured, validated, reconciled, and approved; key controls; error handling; change control; responsible teams/roles.\n"
            "Records & Retention: Where records are stored (systems/locations), retention time, backup/archiving.\n"
            "Reporting Setup: List State pairs operated (high-level); note major Aerodrome pairs if included; any exclusions or special cases.\n"
            "Data Gaps: Whether gaps occurred; % of data affected; replacement method(s) (e.g., QAR, CERT); justification.\n"
            "Emission Reductions / SAF: Any SAF claims, book-and-claim, or emission reductions claimed; documentation required.\n\n"
            "JSON shape (include ALL keys; if unknown use null):\n"
            "{\n"
            "  \"metadata\": { \"document_name\": string, \"mime_type\": string, \"extracted_at\": string, \"generator_version\": string },\n"
            "  \"operator\": { \"name\": {\"value\": string|null, \"__meta\": {\"provenance\": string, \"confidence\": number}}, \"address\": {\"lines\": [string], \"city\": string|null, \"region\": string|null, \"postal_code\": string|null, \"country\": string|null}, \"contacts\": [object], \"aoc\": {\"code\": string|null, \"issued_at\": string|null, \"expires_at\": string|null, \"authority\": {\"name\": string|null, \"address\": object}, \"scope\": [string]} , \"group_structure\": {\"parent_subsidiary_single_entity\": boolean|null, \"subsidiaries\": [{\"name\": string, \"aircraft_identification_method\": string}] }, \"state_of_operator\": string|null },\n"
            "  \"flight_attribution\": { \"icao_designator\": string|null, \"iata_code\": string|null, \"registration_marks\": [string], \"responsibility_under_corsia\": string|null, \"additional_info\": string|null },\n"
            "  \"scope_period\": { \"reporting_year\": string|null, \"coverage\": string|null, \"aggregation_level\": string|null, \"international_only\": boolean|null },\n"
            "  \"activities\": { \"description\": string|null, \"leasing_arrangements\": [string], \"operation_types\": [string], \"geographic_scope\": [string], \"geographical_presence\": string|null },\n"
            "  \"fleet_fuel\": { \"aircraft_list\": [{\"registration_mark\": string, \"type\": string|null, \"count\": number|null, \"notes\": string|null}], \"fuel_types\": [string], \"fuel_suppliers\": [string] },\n"
            "  \"monitoring_methods\": { \"method_a\": object, \"method_b\": object, \"block_off_on\": object, \"fuel_uplift\": object, \"fuel_allocation_block_hour\": object, \"periods_of_applicability\": [string], \"primary_data_sources\": [string], \"sub_fleet_applicability\": object },\n"
            "  \"calculation_inputs\": { \"fuel_density_values\": string|null, \"fuel_density_sources\": string|null, \"cert_usage\": boolean|null, \"cert_inputs\": string|null, \"emission_factors\": string|null },\n"
            "  \"data_management\": { \"data_flow\": string|null, \"controls\": string|null, \"risk_analysis\": string|null, \"data_gaps\": string|null, \"qa_qc_controls\": string|null, \"data_validation_procedures\": string|null, \"change_control\": string|null, \"responsible_teams\": [string] },\n"
            "  \"records_retention\": { \"storage_systems\": [string], \"storage_locations\": [string], \"retention_time\": string|null, \"backup_archiving\": string|null },\n"
            "  \"reporting_setup\": { \"state_pairs_operated\": [string], \"aerodrome_pairs\": [string], \"exclusions\": [string], \"special_cases\": [string] },\n"
            "  \"data_gaps\": { \"gaps_occurred\": boolean|null, \"data_affected_percentage\": number|null, \"replacement_methods\": [string], \"justification\": string|null },\n"
            "  \"emission_reductions_saf\": { \"saf_claims\": [string], \"book_and_claim\": boolean|null, \"emission_reductions_claimed\": [string], \"documentation_required\": [string] },\n"
            "  \"provenance_index\": [{ \"id\": string, \"location\": string, \"excerpt\": string }]\n"
            "}\n"
        )

    def _build_extraction_schema(self) -> Dict[str, Any]:
        """Strict JSON Schema for structured outputs with nested CORSIA sections."""
        return {
            "name": "corsia_monitoring_plan",
            "schema": {
                "type": "object",
                "properties": {
                    "metadata": {
                        "type": "object",
                        "properties": {
                            "document_name": {"type": "string"},
                            "mime_type": {"type": "string"},
                            "extracted_at": {"type": "string"},
                            "generator_version": {"type": "string"}
                        },
                        "required": ["document_name", "mime_type", "extracted_at", "generator_version"],
                        "additionalProperties": False
                    },
                    "operator": {
                        "type": "object",
                        "properties": {
                            "name": {
                                "type": "object",
                                "properties": {
                                    "value": {"type": ["string", "null"]},
                                    "__meta": {
                                        "type": "object",
                                        "properties": {
                                            "provenance": {"type": "string"},
                                            "confidence": {"type": "number"}
                                        },
                                        "required": ["provenance", "confidence"],
                                        "additionalProperties": False
                                    }
                                },
                                "required": ["value", "__meta"],
                                "additionalProperties": False
                            },
                            "address": {
                                "type": "object",
                                "properties": {
                                    "lines": {"type": "array", "items": {"type": "string"}},
                                    "city": {"type": ["string", "null"]},
                                    "region": {"type": ["string", "null"]},
                                    "postal_code": {"type": ["string", "null"]},
                                    "country": {"type": ["string", "null"]}
                                },
                                "required": ["lines", "city", "region", "postal_code", "country"],
                                "additionalProperties": False
                            },
                            "contacts": {
                                "type": "array",
                                "items": {
                                    "type": "object",
                                    "properties": {
                                        "name": {"type": ["string", "null"]},
                                        "title": {"type": ["string", "null"]},
                                        "email": {"type": ["string", "null"]},
                                        "phone": {"type": ["string", "null"]}
                                    },
                                    "required": ["name", "title", "email", "phone"],
                                    "additionalProperties": False
                                }
                            },
                            "aoc": {
                                "type": "object",
                                "properties": {
                                    "code": {"type": ["string", "null"]},
                                    "issued_at": {"type": ["string", "null"]},
                                    "expires_at": {"type": ["string", "null"]},
                                    "authority": {
                                        "type": "object",
                                        "properties": {
                                            "name": {"type": ["string", "null"]},
                                            "address": {
                                                "type": "object",
                                                "properties": {
                                                    "lines": {"type": "array", "items": {"type": "string"}},
                                                    "city": {"type": ["string", "null"]},
                                                    "region": {"type": ["string", "null"]},
                                                    "postal_code": {"type": ["string", "null"]},
                                                    "country": {"type": ["string", "null"]}
                                                },
                                                "required": ["lines", "city", "region", "postal_code", "country"],
                                                "additionalProperties": False
                                            }
                                        },
                                        "required": ["name", "address"],
                                        "additionalProperties": False
                                    },
                                    "scope": {"type": "array", "items": {"type": "string"}}
                                },
                                "required": ["code", "issued_at", "expires_at", "authority", "scope"],
                                "additionalProperties": False
                            },
                            "group_structure": {
                                "type": "object",
                                "properties": {
                                    "parent_subsidiary_single_entity": {"type": ["boolean", "null"]},
                                    "subsidiaries": {
                                        "type": "array",
                                        "items": {
                                            "type": "object",
                                            "properties": {
                                                "name": {"type": "string"},
                                                "aircraft_identification_method": {"type": "string"}
                                            },
                                            "required": ["name", "aircraft_identification_method"],
                                            "additionalProperties": False
                                        }
                                    }
                                },
                                "required": ["parent_subsidiary_single_entity", "subsidiaries"],
                                "additionalProperties": False
                            },
                            "state_of_operator": {"type": ["string", "null"]}
                        },
                        "required": ["name", "address", "contacts", "aoc", "group_structure", "state_of_operator"],
                        "additionalProperties": False
                    },
                    "flight_attribution": {
                        "type": "object",
                        "properties": {
                            "icao_designator": {"type": ["string", "null"]},
                            "iata_code": {"type": ["string", "null"]},
                            "registration_marks": {"type": "array", "items": {"type": "string"}},
                            "responsibility_under_corsia": {"type": ["string", "null"]},
                            "additional_info": {"type": ["string", "null"]}
                        },
                        "required": ["icao_designator", "iata_code", "registration_marks", "responsibility_under_corsia", "additional_info"],
                        "additionalProperties": False
                    },
                    "scope_period": {
                        "type": "object",
                        "properties": {
                            "reporting_year": {"type": ["string", "null"]},
                            "coverage": {"type": ["string", "null"]},
                            "aggregation_level": {"type": ["string", "null"]},
                            "international_only": {"type": ["boolean", "null"]}
                        },
                        "required": ["reporting_year", "coverage", "aggregation_level", "international_only"],
                        "additionalProperties": False
                    },
                    "activities": {
                        "type": "object",
                        "properties": {
                            "description": {"type": ["string", "null"]},
                            "leasing_arrangements": {"type": "array", "items": {"type": "string"}},
                            "operation_types": {"type": "array", "items": {"type": "string"}},
                            "geographic_scope": {"type": "array", "items": {"type": "string"}},
                            "geographical_presence": {"type": ["string", "null"]}
                        },
                        "required": ["description", "leasing_arrangements", "operation_types", "geographic_scope", "geographical_presence"],
                        "additionalProperties": False
                    },
                    "fleet_fuel": {
                        "type": "object",
                        "properties": {
                            "aircraft_list": {
                                "type": "array",
                                "items": {
                                    "type": "object",
                                    "properties": {
                                        "registration_mark": {"type": "string"},
                                        "type": {"type": ["string", "null"]},
                                        "count": {"type": ["number", "null"]},
                                        "notes": {"type": ["string", "null"]}
                                    },
                                    "required": ["registration_mark", "type", "count", "notes"],
                                    "additionalProperties": False
                                }
                            },
                            "fuel_types": {"type": "array", "items": {"type": "string"}},
                            "fuel_suppliers": {"type": "array", "items": {"type": "string"}}
                        },
                        "required": ["aircraft_list", "fuel_types", "fuel_suppliers"],
                        "additionalProperties": False
                    },
                    "monitoring_methods": {
                        "type": "object",
                        "properties": {
                            "method_a": {"type": "object"},
                            "method_b": {"type": "object"},
                            "block_off_on": {"type": "object"},
                            "fuel_uplift": {"type": "object"},
                            "fuel_allocation_block_hour": {"type": "object"},
                            "periods_of_applicability": {"type": "array", "items": {"type": "string"}},
                            "primary_data_sources": {"type": "array", "items": {"type": "string"}},
                            "sub_fleet_applicability": {"type": "object"}
                        },
                        "required": ["method_a", "method_b", "block_off_on", "fuel_uplift", "fuel_allocation_block_hour", "periods_of_applicability", "primary_data_sources", "sub_fleet_applicability"],
                                "additionalProperties": False
                            },
                    "calculation_inputs": {
                        "type": "object",
                        "properties": {
                            "fuel_density_values": {"type": ["string", "null"]},
                            "fuel_density_sources": {"type": ["string", "null"]},
                            "cert_usage": {"type": ["boolean", "null"]},
                            "cert_inputs": {"type": ["string", "null"]},
                            "emission_factors": {"type": ["string", "null"]}
                        },
                        "required": ["fuel_density_values", "fuel_density_sources", "cert_usage", "cert_inputs", "emission_factors"],
                        "additionalProperties": False
                    },
                    "data_management": {
                        "type": "object",
                        "properties": {
                            "data_flow": {"type": ["string", "null"]},
                            "controls": {"type": ["string", "null"]},
                            "risk_analysis": {"type": ["string", "null"]},
                            "data_gaps": {"type": ["string", "null"]},
                            "qa_qc_controls": {"type": ["string", "null"]},
                            "data_validation_procedures": {"type": ["string", "null"]},
                            "change_control": {"type": ["string", "null"]},
                            "responsible_teams": {"type": "array", "items": {"type": "string"}}
                        },
                        "required": ["data_flow", "controls", "risk_analysis", "data_gaps", "qa_qc_controls", "data_validation_procedures", "change_control", "responsible_teams"],
                        "additionalProperties": False
                    },
                    "records_retention": {
                        "type": "object",
                        "properties": {
                            "storage_systems": {"type": "array", "items": {"type": "string"}},
                            "storage_locations": {"type": "array", "items": {"type": "string"}},
                            "retention_time": {"type": ["string", "null"]},
                            "backup_archiving": {"type": ["string", "null"]}
                        },
                        "required": ["storage_systems", "storage_locations", "retention_time", "backup_archiving"],
                        "additionalProperties": False
                    },
                    "reporting_setup": {
                        "type": "object",
                        "properties": {
                            "state_pairs_operated": {"type": "array", "items": {"type": "string"}},
                            "aerodrome_pairs": {"type": "array", "items": {"type": "string"}},
                            "exclusions": {"type": "array", "items": {"type": "string"}},
                            "special_cases": {"type": "array", "items": {"type": "string"}}
                        },
                        "required": ["state_pairs_operated", "aerodrome_pairs", "exclusions", "special_cases"],
                        "additionalProperties": False
                    },
                    "data_gaps": {
                        "type": "object",
                        "properties": {
                            "gaps_occurred": {"type": ["boolean", "null"]},
                            "data_affected_percentage": {"type": ["number", "null"]},
                            "replacement_methods": {"type": "array", "items": {"type": "string"}},
                            "justification": {"type": ["string", "null"]}
                        },
                        "required": ["gaps_occurred", "data_affected_percentage", "replacement_methods", "justification"],
                        "additionalProperties": False
                    },
                    "emission_reductions_saf": {
                        "type": "object",
                        "properties": {
                            "saf_claims": {"type": "array", "items": {"type": "string"}},
                            "book_and_claim": {"type": ["boolean", "null"]},
                            "emission_reductions_claimed": {"type": "array", "items": {"type": "string"}},
                            "documentation_required": {"type": "array", "items": {"type": "string"}}
                        },
                        "required": ["saf_claims", "book_and_claim", "emission_reductions_claimed", "documentation_required"],
                        "additionalProperties": False
                    },
                    "provenance_index": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "id": {"type": "string"},
                                "location": {"type": "string"},
                                "excerpt": {"type": "string"}
                            },
                            "required": ["id", "location", "excerpt"],
                            "additionalProperties": False
                        }
                    }
                },
                "required": [
                    "metadata",
                    "operator",
                    "flight_attribution",
                    "scope_period",
                    "activities",
                    "fleet_fuel",
                    "monitoring_methods",
                    "calculation_inputs",
                    "data_management",
                    "records_retention",
                    "reporting_setup",
                    "data_gaps",
                    "emission_reductions_saf",
                    "provenance_index"
                ],
                "additionalProperties": False
            },
            "strict": True
        }

    def _process_extracted_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Attach extraction metadata and ensure minimal structure without enforcing rigid keys."""
        if not isinstance(data, dict):
            data = {"raw": data}

        meta = data.get('metadata', {}) if isinstance(data.get('metadata'), dict) else {}
        meta.setdefault('extracted_at', pd.Timestamp.now().isoformat())
        meta.setdefault('generator_version', 'v1-flex')
        data['metadata'] = meta

        data['extraction_metadata'] = {
            'model': self.model,
            'reasoning_effort': 'high'
        }

        return data
